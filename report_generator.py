"""
report_generator.py
-------------------
Convert ReportDocument content into physical files:
  - Markdown (.md)
  - PDF (.pdf) via ReportLab
  - DOCX (.docx) via python-docx

Files are written under settings.output_dir (default: output/).
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Union

from config import settings
from models import FileOutputs, ReportDocument
from utils import (
    build_output_basename,
    ensure_dir,
    get_logger,
    log_error,
    log_event,
    slugify,
)

logger = get_logger(__name__)


# =============================================================================
# Public API
# =============================================================================


def generate_report_files(
    document: ReportDocument,
    formats: Optional[Sequence[str]] = None,
    output_dir: Optional[Union[str, Path]] = None,
    request_id: Optional[str] = None,
) -> FileOutputs:
    """
    Write report files for the requested formats.

    Args:
        document: Final report content from Report Writer.
        formats: Subset of md/pdf/docx (defaults to config).
        output_dir: Destination directory.
        request_id: Optional id embedded in filenames.

    Returns:
        FileOutputs with absolute/relative paths that were created.
    """
    if document is None:
        raise ValueError("ReportDocument is required.")

    fmt_list = _normalize_formats(formats or settings.output_formats_list)
    out_dir = ensure_dir(output_dir or settings.output_dir)
    basename = build_output_basename(document.title or "report", request_id=request_id)

    log_event(
        "report_export_started",
        title=document.title,
        formats=fmt_list,
        output_dir=str(out_dir),
        basename=basename,
    )

    paths: Dict[str, Optional[str]] = {"markdown": None, "pdf": None, "docx": None}
    errors: List[str] = []

    if "md" in fmt_list or "markdown" in fmt_list:
        try:
            path = write_markdown(document, out_dir / f"{basename}.md")
            paths["markdown"] = str(path)
        except Exception as exc:  # noqa: BLE001
            log_error("report_export_md_failed", exc)
            errors.append(f"markdown: {exc}")

    if "pdf" in fmt_list:
        try:
            path = write_pdf(document, out_dir / f"{basename}.pdf")
            paths["pdf"] = str(path)
        except Exception as exc:  # noqa: BLE001
            log_error("report_export_pdf_failed", exc)
            errors.append(f"pdf: {exc}")

    if "docx" in fmt_list:
        try:
            path = write_docx(document, out_dir / f"{basename}.docx")
            paths["docx"] = str(path)
        except Exception as exc:  # noqa: BLE001
            log_error("report_export_docx_failed", exc)
            errors.append(f"docx: {exc}")

    if errors and not any(paths.values()):
        raise RuntimeError("Failed to export report in any format: " + "; ".join(errors))

    log_event(
        "report_export_completed",
        title=document.title,
        markdown=paths["markdown"],
        pdf=paths["pdf"],
        docx=paths["docx"],
        partial_errors=errors or None,
    )

    return FileOutputs(
        markdown=paths["markdown"],
        pdf=paths["pdf"],
        docx=paths["docx"],
    )


# =============================================================================
# Markdown
# =============================================================================


def write_markdown(document: ReportDocument, path: Union[str, Path]) -> Path:
    """Write full markdown content to disk."""
    target = Path(path)
    ensure_dir(target.parent)
    content = document.content_markdown or ""
    # Ensure trailing newline for POSIX friendliness
    if content and not content.endswith("\n"):
        content += "\n"
    target.write_text(content, encoding="utf-8")
    return target


# =============================================================================
# PDF (ReportLab) — lightweight Markdown → flowables
# =============================================================================


def write_pdf(document: ReportDocument, path: Union[str, Path]) -> Path:
    """
    Render a readable PDF from Markdown-ish content.
    Supports: H1/H2/H3, paragraphs, bullet lists, simple numbered lists.
    """
    try:
        from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            ListFlowable,
            ListItem,
            PageBreak,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )
    except ImportError as exc:
        raise RuntimeError(
            "reportlab is required for PDF export. pip install reportlab"
        ) from exc

    target = Path(path)
    ensure_dir(target.parent)

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ReportTitle",
            parent=styles["Heading1"],
            fontSize=18,
            spaceAfter=16,
            leading=22,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportH2",
            parent=styles["Heading2"],
            fontSize=14,
            spaceBefore=14,
            spaceAfter=8,
            leading=18,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportH3",
            parent=styles["Heading3"],
            fontSize=12,
            spaceBefore=10,
            spaceAfter=6,
            leading=15,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportBody",
            parent=styles["BodyText"],
            fontSize=10.5,
            leading=14,
            alignment=TA_JUSTIFY,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportBullet",
            parent=styles["BodyText"],
            fontSize=10.5,
            leading=14,
            leftIndent=12,
            spaceAfter=3,
        )
    )

    doc = SimpleDocTemplate(
        str(target),
        pagesize=A4,
        leftMargin=0.85 * inch,
        rightMargin=0.85 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        title=document.title or "Report",
    )

    story: List = []
    md = document.content_markdown or f"# {document.title or 'Report'}\n"

    for block in _iter_md_blocks(md):
        btype = block["type"]
        text = block["text"]

        if btype == "h1":
            story.append(Paragraph(_escape_rl(text), styles["ReportTitle"]))
            story.append(Spacer(1, 6))
        elif btype == "h2":
            story.append(Paragraph(_escape_rl(text), styles["ReportH2"]))
        elif btype == "h3":
            story.append(Paragraph(_escape_rl(text), styles["ReportH3"]))
        elif btype == "bullet":
            story.append(Paragraph(f"• {_escape_rl(text)}", styles["ReportBullet"]))
        elif btype == "number":
            num = block.get("num", "1")
            story.append(
                Paragraph(f"{num}. {_escape_rl(text)}", styles["ReportBullet"])
            )
        elif btype == "blank":
            story.append(Spacer(1, 6))
        else:
            if text.strip():
                story.append(Paragraph(_escape_rl(text), styles["ReportBody"]))

    if not story:
        story.append(Paragraph(_escape_rl(document.title or "Report"), styles["ReportTitle"]))

    doc.build(story)
    return target


# =============================================================================
# DOCX (python-docx)
# =============================================================================


def write_docx(document: ReportDocument, path: Union[str, Path]) -> Path:
    """Render a DOCX file from Markdown-ish content."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for DOCX export. pip install python-docx"
        ) from exc

    target = Path(path)
    ensure_dir(target.parent)

    docx = Document()
    # Sensible default font
    style = docx.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    md = document.content_markdown or f"# {document.title or 'Report'}\n"

    for block in _iter_md_blocks(md):
        btype = block["type"]
        text = block["text"]

        if btype == "h1":
            docx.add_heading(text, level=1)
        elif btype == "h2":
            docx.add_heading(text, level=2)
        elif btype == "h3":
            docx.add_heading(text, level=3)
        elif btype == "bullet":
            docx.add_paragraph(text, style="List Bullet")
        elif btype == "number":
            docx.add_paragraph(text, style="List Number")
        elif btype == "blank":
            continue
        else:
            if text.strip():
                # Basic **bold** support
                p = docx.add_paragraph()
                _add_runs_with_bold(p, text)

    if not docx.paragraphs:
        docx.add_heading(document.title or "Report", level=1)

    docx.save(str(target))
    return target


# =============================================================================
# Markdown helpers
# =============================================================================


_H1_RE = re.compile(r"^#\s+(.*)$")
_H2_RE = re.compile(r"^##\s+(.*)$")
_H3_RE = re.compile(r"^###\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*+]\s+(.*)$")
_NUMBER_RE = re.compile(r"^(\d+)[.)]\s+(.*)$")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _iter_md_blocks(markdown: str) -> Iterable[dict]:
    """
    Very small Markdown block parser for export.
    Not a full CommonMark implementation — good enough for agent reports.
    """
    lines = (markdown or "").replace("\r\n", "\n").split("\n")
    paragraph_buf: List[str] = []

    def flush_para():
        nonlocal paragraph_buf
        if paragraph_buf:
            yield {
                "type": "para",
                "text": " ".join(paragraph_buf).strip(),
            }
            paragraph_buf = []

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            yield from flush_para()
            yield {"type": "blank", "text": ""}
            continue

        # Skip fenced code markers for export simplicity
        if line.strip().startswith("```"):
            yield from flush_para()
            continue

        m = _H1_RE.match(line)
        if m:
            yield from flush_para()
            yield {"type": "h1", "text": m.group(1).strip()}
            continue
        m = _H2_RE.match(line)
        if m:
            yield from flush_para()
            yield {"type": "h2", "text": m.group(1).strip()}
            continue
        m = _H3_RE.match(line)
        if m:
            yield from flush_para()
            yield {"type": "h3", "text": m.group(1).strip()}
            continue
        m = _BULLET_RE.match(line.strip())
        if m:
            yield from flush_para()
            yield {"type": "bullet", "text": m.group(1).strip()}
            continue
        m = _NUMBER_RE.match(line.strip())
        if m:
            yield from flush_para()
            yield {"type": "number", "num": m.group(1), "text": m.group(2).strip()}
            continue

        paragraph_buf.append(line.strip())

    yield from flush_para()


def _escape_rl(text: str) -> str:
    """Escape/convert light Markdown for ReportLab Paragraph (XML-like)."""
    if not text:
        return ""
    # Escape XML special chars first
    out = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    # **bold** → <b>bold</b>
    out = _BOLD_RE.sub(r"<b>\1</b>", out)
    # *italic* (simple)
    out = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<i>\1</i>", out)
    # `code`
    out = re.sub(r"`([^`]+)`", r"<font face='Courier'>\1</font>", out)
    # [text](url) → text (url)
    out = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", out)
    return out


def _add_runs_with_bold(paragraph, text: str) -> None:
    """Add runs to a python-docx paragraph, honoring **bold** segments."""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
    if pos == 0 and not paragraph.runs:
        paragraph.add_run(text)


def _normalize_formats(formats: Sequence[str]) -> List[str]:
    allowed = {"md", "markdown", "pdf", "docx"}
    cleaned: List[str] = []
    for f in formats:
        key = (f or "").strip().lower()
        if not key:
            continue
        if key not in allowed:
            logger.warning("Unsupported output format ignored: %s", key)
            continue
        if key == "markdown":
            key = "md"
        if key not in cleaned:
            cleaned.append(key)
    if not cleaned:
        cleaned = ["md"]
    return cleaned